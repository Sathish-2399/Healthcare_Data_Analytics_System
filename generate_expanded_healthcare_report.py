from collections import Counter
from pathlib import Path
from statistics import mean

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from model import accuracy_score, classification_report, model, vectorizer, X, y
from sklearn.model_selection import train_test_split


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_FILE = BASE_DIR / "MCP_Healthcare_Report_Expanded.docx"
DATA_FILE = BASE_DIR / "Healthcare_Transformed.csv"
RAW_DATA_FILE = BASE_DIR / "Healthcare_data.csv"
IMAGE_DIR = BASE_DIR / "ppt_media"


def set_document_language(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(12)


def set_page_layout(document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(document, text, *, bold=False, italic=False, size=12, align="justify",
                  first_line=0.5, spacing_after=6, line_spacing=1.5):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(first_line)
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(document, text, *, level=1, centered=False, page_break=False):
    if page_break:
        document.add_page_break()
    size_map = {1: 16, 2: 14, 3: 13}
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=size_map.get(level, 12), bold=True)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.line_spacing = 1.4
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def add_simple_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = str(left)
        row[1].text = str(right)
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.2
                for run in paragraph.runs:
                    set_run_font(run, size=12)
    return table


def add_three_column_table(document, headers, data_rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for index, title in enumerate(headers):
        hdr[index].text = title
        for paragraph in hdr[index].paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True)
    for values in data_rows:
        row = table.add_row().cells
        for index, value in enumerate(values):
            row[index].text = str(value)
            for paragraph in row[index].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=12)
    return table


def add_figure(document, image_path, caption, width_inches=5.8):
    if image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        add_paragraph(document, caption, align="center", bold=False, italic=False, size=11, first_line=0)
    else:
        add_paragraph(document, f"{caption} (image not available)", align="center", size=11, first_line=0)


def build_stats():
    df = pd.read_csv(DATA_FILE)
    raw_df = pd.read_csv(RAW_DATA_FILE)

    X_vec = vectorizer.transform(X)
    X_train, X_test, y_train, y_test = train_test_split(
        X_vec, y, test_size=0.2, random_state=42
    )
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    disease_counts = df["disease"].value_counts().to_dict()
    gender_counts = df["gender"].value_counts().to_dict()
    age_group_counts = df["age_group"].value_counts().to_dict()
    symptom_lengths = [len(str(x).split(",")) for x in df["symptoms"]]
    most_common_symptoms = Counter(
        symptom.strip()
        for row in df["symptoms"]
        for symptom in str(row).split(",")
    ).most_common(10)

    return {
        "df": df,
        "raw_rows": len(raw_df),
        "clean_rows": len(df),
        "duplicates_removed": int(raw_df.duplicated().sum()),
        "age_min": int(df["age"].min()),
        "age_max": int(df["age"].max()),
        "age_mean": round(df["age"].mean(), 2),
        "disease_counts": disease_counts,
        "gender_counts": gender_counts,
        "age_group_counts": age_group_counts,
        "symptom_mean": round(mean(symptom_lengths), 2),
        "top_symptoms": most_common_symptoms,
        "accuracy": round(accuracy_score(y_test, y_pred) * 100, 2),
        "report": classification_report(y_test, y_pred, digits=2),
    }


def build_cover_page(document):
    add_paragraph(
        document,
        "HEALTHCARE DATA ANALYTICS SYSTEM",
        bold=True,
        size=15,
        align="center",
        first_line=0,
        spacing_after=8,
        line_spacing=1.3,
    )
    add_paragraph(
        document,
        "23CS1ME - MINI-CAPSTONE PROJECT REPORT",
        bold=True,
        size=14,
        align="center",
        first_line=0,
        spacing_after=10,
        line_spacing=1.3,
    )
    add_paragraph(document, "Submitted by", italic=True, size=14, align="center", first_line=0)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line1 = para.add_run("SATHISH B")
    set_run_font(line1, size=14, bold=True)
    line2 = para.add_run("\n2312099")
    set_run_font(line2, size=13, bold=True)
    add_paragraph(document, "In partial fulfillment for the award of the degree", align="center", first_line=0)
    add_paragraph(document, "of", align="center", first_line=0)
    add_paragraph(document, "BACHELOR OF ENGINEERING", bold=True, size=14, align="center", first_line=0)
    add_paragraph(document, "in", align="center", first_line=0)
    add_paragraph(document, "COMPUTER SCIENCE AND ENGINEERING", bold=True, size=14, align="center", first_line=0)
    add_paragraph(document, "NATIONAL ENGINEERING COLLEGE", bold=True, size=14, align="center", first_line=0)
    add_paragraph(
        document,
        "(An Autonomous Institution affiliated to Anna University, Chennai)",
        align="center",
        first_line=0,
    )
    add_paragraph(document, "K.R. Nagar, Kovilpatti - 628503", align="center", first_line=0)
    add_paragraph(document, "APRIL - 2026", bold=True, align="center", first_line=0)


def build_certificate_page(document):
    add_heading(document, "BONAFIDE CERTIFICATE", level=1, centered=True, page_break=True)
    add_paragraph(
        document,
        'This is to certify that the project report entitled "Healthcare Data Analytics System" is the bonafide work '
        "of SATHISH B (2312099), Department of Computer Science and Engineering, National Engineering College, "
        "Kovilpatti, carried out by him during the Mini-Capstone Project course in partial fulfillment of the "
        "requirements for the award of the Degree of Bachelor of Engineering in Computer Science and Engineering.",
    )
    add_paragraph(
        document,
        "The work embodied in this report is original to the extent that it reflects the design, development, data "
        "preparation, experimentation, documentation, and presentation completed for the academic purpose of the "
        "mini-capstone evaluation. The report has been prepared under the guidance of the course instructor and "
        "domain mentor and has not been submitted elsewhere for the award of any other degree or diploma.",
    )
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Project Coordinator"
    table.rows[0].cells[1].text = "Head of the Department"
    table.rows[1].cells[0].text = "Ms. P. Priyadharshini\nAssistant Professor / CSE"
    table.rows[1].cells[1].text = "Department of CSE\nNational Engineering College"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=True if row == table.rows[0] else False)
    add_paragraph(
        document,
        "Submitted to the Mini-Capstone Project Viva-Voce examination held at National Engineering College, "
        "Kovilpatti, during April 2026.",
        first_line=0,
    )


def build_acknowledgement(document):
    add_heading(document, "ACKNOWLEDGEMENT", level=1, centered=True, page_break=True)
    paragraphs = [
        "I express my sincere gratitude to the management and principal of National Engineering College for creating "
        "an environment that encourages innovation, disciplined learning, and practical project development. The "
        "facilities, support systems, and institutional motivation provided during the mini-capstone semester helped "
        "me complete this report and the associated software components with confidence.",
        "I record my heartfelt thanks to the faculty members of the Department of Computer Science and Engineering "
        "for their technical support, academic guidance, and valuable suggestions at every stage of the work. Their "
        "feedback helped shape the project from a basic idea into a structured healthcare analytics solution that "
        "includes data preparation, prediction logic, and report-level documentation.",
        "I extend special thanks to Ms. P. Priyadharshini, Assistant Professor / CSE and project coordinator, for "
        "her steady encouragement, technical inputs, and academic mentoring. Her suggestions on project scope, "
        "clarity of explanation, and implementation discipline significantly improved the final quality of the work.",
        "I also thank my classmates, friends, and family members for their support during coding, data inspection, "
        "dashboard preparation, presentation planning, and report compilation. Their encouragement helped me "
        "maintain continuity and complete the project in a timely manner.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)


def build_abstract(document, stats):
    add_heading(document, "ABSTRACT", level=1, centered=True, page_break=True)
    paragraphs = [
        "Healthcare institutions continuously generate patient information in the form of demographics, symptoms, "
        "clinical notes, and disease labels. When these records are left unorganized, the valuable relationships "
        "between symptoms and disease outcomes remain hidden, making quick interpretation difficult. The present "
        "project proposes a Healthcare Data Analytics System that combines data cleaning, machine learning-based "
        "prediction, and web-based visualization so that patient information can be transformed into actionable "
        "insight for preliminary decision support.",
        f"The transformed dataset used in this project contains {stats['clean_rows']} valid patient records produced "
        f"from a raw source file of {stats['raw_rows']} rows. The preprocessing workflow removes duplicate records, "
        "checks age validity, confirms the consistency of symptom counts, fills missing values, standardizes text, "
        "and derives an age-group attribute. This step improves data quality and ensures that the downstream model "
        "receives reliable inputs for training and inference.",
        "For predictive analytics, the project uses a Multinomial Naive Bayes classifier trained on a combined "
        "text representation of age, gender, age group, and symptoms. CountVectorizer converts the textual input "
        "into numeric feature vectors, after which the classifier learns the symptom-disease associations embedded "
        "within the dataset. On the current test split, the implemented model achieved an accuracy of "
        f"{stats['accuracy']}%, demonstrating that the feature engineering strategy is effective for the available "
        "healthcare records.",
        "The application layer is implemented using Flask. The user enters the patient name, age, gender, and "
        "symptoms through an HTML form, and the application returns the top three probable diseases as HIGH, MEDIUM, "
        "and LOW ranked outcomes. Alongside this predictive component, a browser-based dashboard reads the "
        "transformed dataset and presents descriptive analytics such as disease distribution, age-group patterns, "
        "gender distribution, and top-disease summaries. The project therefore demonstrates the practical value of "
        "combining predictive analytics with descriptive analytics in a single academic mini-capstone workflow.",
        "This report has been expanded in a detailed academic format covering the project background, problem "
        "definition, literature survey, design thinking, methodology, architecture, implementation, evaluation, "
        "result interpretation, limitations, future scope, references, and appendix materials. The report aims to "
        "support final submission, viva preparation, and project review by presenting the work in a clear, complete, "
        "and presentation-ready manner.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)


def build_contents(document):
    add_heading(document, "TABLE OF CONTENTS", level=1, centered=True, page_break=True)
    contents = [
        "Certificate",
        "Acknowledgement",
        "Abstract",
        "Chapter 1  Introduction",
        "Chapter 2  Literature Survey",
        "Chapter 3  Design Thinking and Problem Framing",
        "Chapter 4  Proposed Methodology and System Design",
        "Chapter 5  Implementation, Analytics, and Results",
        "Chapter 6  Discussion, Limitations, and Future Scope",
        "Chapter 7  Conclusion",
        "Appendix I   Sample Inputs, Outputs, and Screens",
        "Appendix II  Code Structure and File Description",
        "References",
    ]
    for item in contents:
        add_number(document, item)

    add_heading(document, "LIST OF FIGURES", level=2, centered=False)
    figure_items = [
        "Fig. 4.1 Overall project workflow",
        "Fig. 4.2 Layered system architecture",
        "Fig. 5.1 Prediction module screen",
        "Fig. 5.2 Analytics dashboard view",
        "Fig. 5.3 Additional dashboard charts",
        "Fig. A1 Screenshot set from project presentation assets",
    ]
    for item in figure_items:
        add_bullet(document, item)

    add_heading(document, "LIST OF TABLES", level=2, centered=False)
    table_items = [
        "Table 4.1 Functional modules",
        "Table 4.2 Software and hardware requirements",
        "Table 5.1 Cleaned dataset summary",
        "Table 5.2 Disease distribution",
        "Table 5.3 Gender and age-group distribution",
        "Table 5.4 Model evaluation summary",
    ]
    for item in table_items:
        add_bullet(document, item)


def chapter_one(document, stats):
    add_heading(document, "CHAPTER 1", level=1, centered=True, page_break=True)
    add_heading(document, "INTRODUCTION", level=1, centered=True)
    paragraphs = [
        "Healthcare data analytics has become one of the most important application areas of computer science "
        "because patient-related data is growing rapidly in digital form. Hospitals, diagnostic centers, clinics, "
        "telemedicine services, and academic repositories continuously record disease labels, patient demographics, "
        "and symptom descriptions. Even though this information is valuable, healthcare workers often need fast and "
        "easy ways to interpret it. When data remains in raw tabular form without proper processing, the clinical "
        "patterns hidden in the records are not easy to understand or communicate.",
        "The project titled Healthcare Data Analytics System addresses this challenge by designing a compact system "
        "that combines preprocessing, disease prediction, and dashboard-based analysis. The goal is not to replace a "
        "doctor or a laboratory diagnosis. Instead, the proposed solution acts as a lightweight decision-support "
        "tool that demonstrates how a cleaned dataset and a simple machine learning pipeline can convert symptom "
        "information into meaningful ranked outcomes. The project is therefore academically strong because it covers "
        "the full path from raw data to user-facing output.",
        "The system begins with a healthcare dataset that includes age, gender, symptoms, symptom count, and disease "
        "class. In its raw form, such data may contain duplicate rows, inconsistent spacing, missing values, and "
        "other quality issues. The project first resolves these issues through structured preprocessing. It then "
        "creates an engineered input string that combines age, gender, age group, and symptoms. That combined text "
        "representation is converted into machine-readable vectors and used to train a Multinomial Naive Bayes "
        "classifier.",
        "The trained model is integrated into a Flask application so that users can interact with the system through "
        "a clean browser interface. Rather than returning just one predicted disease, the application reports the top "
        "three disease possibilities and labels them as HIGH, MEDIUM, and LOW probability outcomes. This design is "
        "useful in healthcare-like scenarios because multiple diseases can share overlapping symptoms. Ranked output "
        "provides a broader and more realistic interpretation than a single-label response.",
        "In addition to prediction, the project includes an analytics dashboard built with HTML, JavaScript, and "
        "Chart.js. The dashboard reads the transformed CSV file and visualizes the distribution of diseases, gender, "
        "and age groups. It also presents summary indicators such as total records, average age, and top disease. "
        "This dual capability makes the project stronger than a simple prediction demo because it demonstrates both "
        "descriptive analytics and predictive analytics using the same project data.",
        f"The current transformed dataset contains {stats['clean_rows']} records, with patient ages ranging from "
        f"{stats['age_min']} to {stats['age_max']} years and an average age of {stats['age_mean']}. The largest "
        f"disease categories are {', '.join(list(stats['disease_counts'].keys())[:5])}. These figures indicate that "
        "the dataset has enough variety to support meaningful experimentation and discussion during academic review.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)

    add_heading(document, "1.1 NEED FOR THE PROJECT", level=2)
    need_paragraphs = [
        "The need for a healthcare analytics project arises from the gap between data availability and data usability. "
        "Modern healthcare systems collect significant amounts of data, but a large portion of it is used only for "
        "record storage rather than pattern discovery. Students and practitioners need examples of compact systems "
        "that show how raw healthcare records can be transformed into practical insight with the help of machine "
        "learning and visualization.",
        "Another need comes from the educational importance of interpretable analytics. Complex deep learning systems "
        "may produce strong results, but they are harder to explain during mini-capstone reviews. A Naive Bayes "
        "workflow, by contrast, is easier to justify, implement, and demonstrate. It offers a good balance between "
        "accuracy, simplicity, and academic clarity.",
        "The project also responds to the need for quick symptom-based support tools. In many introductory healthcare "
        "informatics settings, users require a fast way to test how combinations of symptoms relate to diseases. "
        "This system supports that need by taking structured input and returning an immediate ranked prediction.",
    ]
    for text in need_paragraphs:
        add_paragraph(document, text)

    add_heading(document, "1.2 OBJECTIVES", level=2)
    objectives = [
        "To clean and transform raw healthcare records into a consistent analytical dataset.",
        "To derive meaningful features such as age group and combined symptom-demographic text input.",
        "To build a machine learning model that predicts probable diseases from patient information.",
        "To create a Flask-based front-end for accepting inputs and displaying top-ranked disease outcomes.",
        "To design a dashboard that visualizes disease, gender, and age-group distributions from the transformed data.",
        "To document the project in a report format suitable for academic review and final submission.",
    ]
    for item in objectives:
        add_bullet(document, item)

    add_heading(document, "1.3 SCOPE OF THE PROJECT", level=2)
    scope = [
        "The scope of the project includes healthcare data cleaning, symptom-based disease classification, and "
        "descriptive dashboard generation using the available dataset. The solution is designed for educational "
        "demonstration, prototype-level analytics, and mini-capstone evaluation. It is suitable for showing how "
        "machine learning can support healthcare interpretation in a controlled academic setting.",
        "The scope deliberately excludes direct clinical validation, electronic medical record integration, and real "
        "patient treatment decisions. The current implementation relies on the structure of the provided dataset and "
        "is optimized for clarity and explainability rather than hospital-scale deployment. Even with those "
        "limitations, the system successfully demonstrates the core pipeline needed for healthcare analytics.",
    ]
    for text in scope:
        add_paragraph(document, text)

    add_heading(document, "1.4 PROBLEM STATEMENT", level=2)
    add_paragraph(
        document,
        "Healthcare datasets often contain hidden relationships among demographics, symptom combinations, and disease "
        "labels, but these relationships are difficult to identify manually. Existing lightweight academic systems "
        "either focus only on static analysis or only on prediction, leaving a gap for an integrated solution that "
        "cleans data, predicts likely diseases, and visualizes population-level patterns. The problem addressed in "
        "this project is how to build a simple, explainable, and accurate healthcare analytics system that transforms "
        "raw patient records into clear predictive and descriptive outputs for academic use.",
    )

    add_heading(document, "1.5 SIGNIFICANCE OF THE WORK", level=2)
    significance = [
        "The work demonstrates that even a lightweight pipeline can achieve strong predictive performance when the "
        "data is well prepared and the feature design is meaningful.",
        "It shows how one transformed dataset can support multiple outputs such as prediction pages, API responses, "
        "and dashboard charts.",
        "It improves project presentation value because the user can explain preprocessing, machine learning, web "
        "development, and visualization within a single report.",
        "It offers a reusable foundation that can be extended later with new models, richer healthcare features, or "
        "persistent database support.",
    ]
    for item in significance:
        add_bullet(document, item)


def chapter_two(document):
    add_heading(document, "CHAPTER 2", level=1, centered=True, page_break=True)
    add_heading(document, "LITERATURE SURVEY", level=1, centered=True)
    sections = {
        "2.1 OVERVIEW OF EXISTING HEALTHCARE ANALYTICS SYSTEMS": [
            "Healthcare analytics systems in the literature generally fall into two broad categories. The first "
            "category focuses on descriptive analytics, where dashboards summarize patient populations, disease rates, "
            "utilization trends, and demographic distributions. The second category focuses on predictive analytics, "
            "where supervised learning algorithms classify diseases or estimate outcomes using historical patient "
            "data. Both categories are important, but many academic prototypes emphasize one while ignoring the other.",
            "Descriptive systems are useful because they help users understand dataset composition, identify dominant "
            "classes, and communicate trends visually. However, they may not answer patient-specific questions. "
            "Predictive systems are useful because they generate direct outcomes from input values, but when they do "
            "not include supporting data context, users may find them difficult to trust or explain. This project "
            "combines both perspectives in a single workflow.",
        ],
        "2.2 RULE-BASED SYMPTOM CHECKERS": [
            "Traditional symptom checker tools often rely on manually crafted rules. For example, if a patient shows "
            "a given combination of fever, cough, and fatigue, the system may map that pattern to a predefined "
            "disease group. Rule-based systems are simple and explainable, but they become difficult to maintain as "
            "the number of symptom combinations grows. They also struggle to represent uncertainty and overlap among "
            "diseases because medical conditions frequently share common symptoms.",
            "The limitation of rule-based systems motivates the use of data-driven methods that learn from examples. "
            "Instead of requiring a human expert to define every condition, supervised models infer statistical "
            "patterns from labeled data. This makes them more flexible in academic and practical analytics settings.",
        ],
        "2.3 MACHINE LEARNING IN DISEASE PREDICTION": [
            "Machine learning is widely used in disease prediction because it can generalize from historical patient "
            "records. Common algorithms include Decision Tree, Random Forest, Support Vector Machine, K-Nearest "
            "Neighbors, Logistic Regression, Naive Bayes, and neural networks. The choice of algorithm depends on "
            "dataset size, feature type, explainability needs, and deployment constraints.",
            "For text-like symptom input, Naive Bayes is often a strong baseline because it handles sparse feature "
            "representations efficiently and produces quick predictions. CountVectorizer-based pipelines are "
            "particularly suitable when symptoms are represented as tokens or short textual phrases. These properties "
            "align well with the scope of the present project, where fast inference and explainable implementation are "
            "important.",
        ],
        "2.4 DATA QUALITY IN HEALTHCARE ANALYTICS": [
            "Data quality plays a decisive role in model performance. In healthcare datasets, common issues include "
            "duplicate records, inconsistent text formatting, missing values, unrealistic age values, and mismatched "
            "symptom metadata. If such issues are ignored, the resulting model may appear accurate during some tests "
            "while actually learning misleading patterns. Therefore, preprocessing is not a minor step but a central "
            "component of healthcare analytics.",
            "The reviewed practices in data science recommend explicit validation and transformation before training. "
            "These include standardizing field names, normalizing strings, validating numeric ranges, checking "
            "consistency constraints, imputing missing values, and deriving new informative features. This project "
            "adopts these principles directly in the preprocessing script.",
        ],
        "2.5 DASHBOARDS FOR ANALYTICAL INTERPRETATION": [
            "Dashboards transform tabular records into visual summaries that are easy to interpret. In healthcare "
            "settings, dashboards often include bar charts for disease counts, pie charts for gender split, trend "
            "lines for patient volume, and cross-tab views comparing disease against age groups or risk levels. They "
            "are useful for presentations because they turn raw data into visible evidence.",
            "The current project uses a dashboard not as a decorative add-on but as an analytical layer that supports "
            "the prediction engine. This decision improves report quality because the user can explain the dataset "
            "composition, identify dominant disease classes, and justify model behavior using the same transformed "
            "source file.",
        ],
        "2.6 RESEARCH GAP AND PROJECT POSITIONING": [
            "From the surveyed approaches, a clear gap appears in the area of compact academic systems that integrate "
            "preprocessing, prediction, and dashboard analytics in a single, easy-to-explain pipeline. Many projects "
            "either contain only theoretical discussion or present only one technical component. The present work "
            "fills that gap by offering a complete but lightweight stack that is appropriate for mini-capstone "
            "evaluation.",
            "The project positions itself as an explainable healthcare analytics prototype. Its strength lies in the "
            "balance it maintains across data quality, model simplicity, interface usability, and report readiness. "
            "That balance makes it a strong educational artifact and a suitable base for future expansion.",
        ],
    }
    for heading, paragraphs in sections.items():
        add_heading(document, heading, level=2)
        for text in paragraphs:
            add_paragraph(document, text)


def chapter_three(document):
    add_heading(document, "CHAPTER 3", level=1, centered=True, page_break=True)
    add_heading(document, "DESIGN THINKING AND PROBLEM FRAMING", level=1, centered=True)

    add_heading(document, "3.1 EMPATHIZE", level=2)
    empathize = [
        "The intended users of the system include students, faculty reviewers, healthcare learners, and anyone who "
        "needs a quick educational understanding of how symptom-based analytics can work. These users are not always "
        "experts in machine learning, so the interface and the report must remain simple, readable, and transparent.",
        "During project framing, it became clear that users want two things at the same time: an immediate answer for "
        "specific symptom input and a broader explanation of how the dataset behaves. This is why the project was "
        "designed with both a prediction form and a dashboard. The first satisfies the need for direct interaction, "
        "while the second satisfies the need for interpretation and trust.",
    ]
    for text in empathize:
        add_paragraph(document, text)
    for item in [
        "Users say they need a system that gives a quick idea of likely diseases from symptom descriptions.",
        "Users think the tool should be simple enough to use during demonstrations without technical confusion.",
        "Users see many healthcare records but few compact examples that show prediction and analytics together.",
        "Users hear that machine learning can help in healthcare, but they want understandable outputs.",
        "Users feel pain when manual inspection of symptoms takes time and gives uncertain conclusions.",
        "Users gain value when the system explains likely outcomes and shows supporting dataset trends.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3.2 DEFINE", level=2)
    define_paragraphs = [
        "The design challenge can be defined as follows: build a healthcare analytics solution that accepts basic "
        "patient attributes, predicts probable diseases accurately, and also presents the dataset in a form that is "
        "easy to explain during project review. The solution must be lightweight, easy to implement using standard "
        "Python tools, and suitable for final-year academic demonstration.",
        "The system should not overload the user with technical detail on the interface. Instead, complexity should "
        "remain inside the preprocessing and modeling logic, while the front-end exposes only the necessary fields "
        "and ranked outputs. The report itself should then provide the technical depth needed for academic assessment.",
    ]
    for text in define_paragraphs:
        add_paragraph(document, text)

    add_heading(document, "3.3 IDEATE", level=2)
    ideation = [
        "Several implementation options were considered during ideation. One option was to build only a chart-based "
        "dashboard. This would have produced descriptive insight but no predictive interaction. A second option was "
        "to build only a machine learning prediction form. This would give direct output but would not help explain "
        "the dataset. The final and most balanced idea was to combine both approaches.",
        "For the learning algorithm, a simple interpretable method was preferred over a heavyweight architecture. "
        "Given that symptoms can be represented as short text and the dataset size is moderate, CountVectorizer with "
        "Multinomial Naive Bayes emerged as a practical solution. It is computationally efficient, academically "
        "defensible, and easy to integrate with a Flask application.",
        "For data storage, CSV was chosen as the primary working format because it supports easy inspection and fits "
        "the project scope. An optional MySQL loading script was also included to show how the transformed data could "
        "be moved into a database system in a future extension or deployment scenario.",
    ]
    for text in ideation:
        add_paragraph(document, text)

    add_heading(document, "3.4 PROTOTYPE", level=2)
    prototype = [
        "The prototype was divided into clear modules. The preprocessing script reads and validates raw data. The "
        "model script trains and evaluates the classifier. The Flask application binds the model to a user form. The "
        "dashboard page renders visual summaries from the transformed dataset. This modular approach simplified "
        "testing and made the project easier to explain chapter by chapter.",
        "The first prototype focused on validating whether age, gender, and symptoms could be merged into a single "
        "feature string. After confirming the feasibility of the approach, the report structure was expanded to "
        "document the full system life cycle, including analytics, architecture, and future extensions.",
    ]
    for text in prototype:
        add_paragraph(document, text)

    add_heading(document, "3.5 TEST AND REFINE", level=2)
    test_refine = [
        "Testing in this project involved two layers. The first layer verified whether preprocessing produced a clean "
        "and consistent dataset. The second layer checked whether the model returned sensible ranked predictions and "
        "whether the dashboard represented aggregate values accurately. Observing strong classification accuracy and "
        "clear dashboard output confirmed that the selected pipeline was suitable for the mini-capstone context.",
        "The final refinement centered on documentation quality. The report was expanded to include deeper academic "
        "discussion, more chapter-wise elaboration, formal front matter, and appendix materials so that the project "
        "presentation would align better with a full-length report format similar to stronger peer submissions.",
    ]
    for text in test_refine:
        add_paragraph(document, text)


def chapter_four(document, stats):
    add_heading(document, "CHAPTER 4", level=1, centered=True, page_break=True)
    add_heading(document, "PROPOSED METHODOLOGY AND SYSTEM DESIGN", level=1, centered=True)

    add_heading(document, "4.1 EXISTING SYSTEM", level=2)
    existing = [
        "Existing lightweight healthcare projects often suffer from fragmentation. Some systems perform basic charting "
        "but do not offer prediction. Others deliver predictions without giving users insight into the underlying "
        "dataset. In addition, many academic prototypes do not emphasize data cleaning, even though that step strongly "
        "influences the reliability of the final output.",
        "Another weakness of the existing approach is the absence of ranking. Systems that output only one disease "
        "label can be misleading because symptom overlap is common. A more practical educational design should show "
        "multiple probable classes so that the user can understand uncertainty and similarity among conditions.",
    ]
    for text in existing:
        add_paragraph(document, text)

    add_heading(document, "4.2 PROPOSED SYSTEM", level=2)
    proposed = [
        "The proposed system begins with raw healthcare data and moves through a disciplined pipeline of cleaning, "
        "validation, feature engineering, model training, prediction, and visualization. The system produces both "
        "record-level outputs and dataset-level insights. This makes it suitable for demonstration, explanation, and "
        "future enhancement.",
        "At the core of the project is a supervised learning approach in which age, gender, age group, and symptoms "
        "are combined into a single text feature. CountVectorizer encodes that feature numerically, and Multinomial "
        "Naive Bayes learns the associations between the encoded input and the disease labels. The model returns class "
        "probabilities, which are then normalized and presented as ranked outcomes to the end user.",
    ]
    for text in proposed:
        add_paragraph(document, text)
    for item in [
        "Input: name, age, gender, and symptom list.",
        "Preprocessing: duplicate removal, age validation, symptom-count verification, missing-value handling, age-group derivation.",
        "Feature engineering: combined string of age, gender, age group, and symptoms.",
        "Modeling: CountVectorizer plus Multinomial Naive Bayes.",
        "Output: top three probable diseases labeled HIGH, MEDIUM, and LOW.",
        "Analytics: summary cards and charts based on the transformed dataset.",
    ]:
        add_bullet(document, item)

    add_heading(document, "4.3 SYSTEM WORKFLOW", level=2)
    workflow = [
        "The workflow starts when the raw healthcare CSV is loaded into the preprocessing script. The script checks "
        "for duplicates, normalizes column names, verifies that age values are realistic, and confirms that the "
        "declared symptom count matches the actual count derived from the symptom string. Missing values are handled "
        "column by column using median or mode strategies depending on data type.",
        "Once the dataset is transformed and saved, the modeling script reads the clean file and constructs the "
        "combined input string. The vectorizer converts tokens to a sparse numerical matrix, and the Naive Bayes "
        "classifier trains on the result. During inference, the same vectorizer transforms new user input, ensuring "
        "consistency between training and prediction phases.",
        "The Flask application connects the prediction function to an HTML form. When the user submits values, the "
        "application derives the age group, invokes the model logic, and renders the ranked response. In parallel, "
        "the analytics dashboard reads the transformed data and generates visual summaries for interpretation.",
    ]
    for text in workflow:
        add_paragraph(document, text)

    add_paragraph(document, "Fig. 4.1 Overall project workflow diagram placeholder", align="center", first_line=0, italic=True)

    add_heading(document, "4.4 SYSTEM ARCHITECTURE", level=2)
    architecture = [
        "The architecture of the proposed project can be described as a layered system with clear separation of "
        "concerns. The data layer contains the CSV-based healthcare records. The preprocessing layer performs quality "
        "control and feature derivation. The machine learning layer transforms symptom-rich inputs into disease "
        "predictions. The presentation layer includes the Flask form interface, API endpoint, and dashboard page.",
        "This layered architecture improves maintainability. Changes in one layer do not require a complete rewrite "
        "of the others. For instance, a future project version could replace Naive Bayes with another classifier while "
        "keeping the same input form and dashboard. Similarly, the storage layer could be moved from CSV to database "
        "without changing the conceptual design of the user-facing modules.",
    ]
    for text in architecture:
        add_paragraph(document, text)
    add_paragraph(document, "Fig. 4.2 Layered system architecture placeholder", align="center", first_line=0, italic=True)

    add_heading(document, "4.5 FUNCTIONAL MODULES", level=2)
    add_three_column_table(
        document,
        ["Module", "Primary Responsibility", "Output"],
        [
            ("Data Preprocessing", "Validate, clean, and transform raw records", "Healthcare_Transformed.csv"),
            ("Feature Engineering", "Combine demographic and symptom fields", "Model input string"),
            ("Prediction Engine", "Train and infer diseases using Naive Bayes", "Ranked disease probabilities"),
            ("Web Application", "Accept patient inputs and show results", "User-facing prediction page"),
            ("Analytics Dashboard", "Visualize aggregate dataset patterns", "Charts and summary cards"),
            ("Database Loader", "Optional upload into MySQL storage", "healthcare_data table"),
        ],
    )
    add_paragraph(
        document,
        "Table 4.1 summarizes the major modules. Each module contributes a distinct part of the system and supports "
        "clean explanation during the final review. The separation of modules also makes the source code easy to read "
        "because each file has a focused responsibility.",
    )

    add_heading(document, "4.6 DATA PREPROCESSING DETAILS", level=2)
    preprocessing = [
        f"The raw dataset originally contains {stats['raw_rows']} rows, while the transformed dataset retains "
        f"{stats['clean_rows']} valid rows. Duplicate rows removed during the process amount to "
        f"{stats['duplicates_removed']}. Although the duplicate count alone does not represent the full value of "
        "preprocessing, it indicates that quality control is active and necessary.",
        "Column names are standardized by trimming spaces, converting text to lowercase, and replacing blank spaces "
        "with underscores. Age values are filtered to retain only realistic patient ages between 1 and 120. A derived "
        "symptom count is computed and matched against the recorded symptom_count field so that structurally "
        "inconsistent rows can be removed. Missing numeric values are filled with the median, while missing "
        "categorical values are filled with the mode.",
        "Finally, the project derives an age_group feature. This step is important because age is not only a number "
        "but also a clinically meaningful category. Children, teenagers, adults, and senior citizens often show "
        "different disease distributions. By including age group in the combined input, the model gains a richer "
        "representation of the patient context.",
    ]
    for text in preprocessing:
        add_paragraph(document, text)

    add_heading(document, "4.7 HARDWARE AND SOFTWARE REQUIREMENTS", level=2)
    add_three_column_table(
        document,
        ["Category", "Specification", "Purpose"],
        [
            ("Processor", "Standard laptop/desktop processor", "Model training and web execution"),
            ("RAM", "4 GB or above", "Data processing and browser usage"),
            ("Operating System", "Windows environment", "Project development and report generation"),
            ("Programming Language", "Python", "Preprocessing, modeling, and backend"),
            ("Libraries", "Pandas, scikit-learn, Flask, SQLAlchemy", "Analytics and deployment support"),
            ("Frontend", "HTML, CSS, JavaScript, Chart.js", "Prediction UI and dashboard"),
            ("Storage", "CSV and optional MySQL", "Dataset persistence"),
        ],
    )

    add_heading(document, "4.8 ALGORITHM CHOICE JUSTIFICATION", level=2)
    algorithm_choice = [
        "Multinomial Naive Bayes was selected because the project uses tokenized text-like symptom input. The "
        "algorithm works well with count-based sparse representations and is computationally light. These properties "
        "are important in educational prototypes where quick execution and clear explanation matter.",
        "A second reason for the choice is interpretability. Even though Naive Bayes makes simplifying assumptions, "
        "its decision process is easier to explain than more complex models. This is beneficial during viva because "
        "the student can clearly justify how the feature vector is built, how training occurs, and why probability-"
        "based ranking is generated.",
        "The algorithm is also easy to integrate with CountVectorizer and Flask. The full pipeline remains compact, "
        "which reduces implementation overhead while preserving strong performance on the current dataset.",
    ]
    for text in algorithm_choice:
        add_paragraph(document, text)

    add_heading(document, "4.9 USE CASE NARRATIVE", level=2)
    use_case = [
        "The primary actor in the system is the end user. The user opens the web application and enters the patient "
        "name, age, gender, and symptom description. Once the form is submitted, the system interprets the age and "
        "converts it into a meaningful age group. This derived value, together with the other user inputs, becomes "
        "the input for the trained disease prediction module.",
        "A secondary use case involves data exploration. Instead of entering one record, the user can view the "
        "dashboard to understand how the transformed dataset is distributed across disease classes, age groups, and "
        "gender categories. This supports educational analysis and helps explain why some disease classes are likely "
        "to appear more often in the prediction output.",
        "Another use case is data persistence. Through datastore.py, the transformed dataset can be loaded into a "
        "MySQL database. This does not change the main behavior of the academic prototype, but it demonstrates how "
        "the project can move toward a more deployment-oriented architecture in later versions.",
    ]
    for text in use_case:
        add_paragraph(document, text)

    add_heading(document, "4.10 STEPWISE ALGORITHM", level=2)
    for item in [
        "Read the raw healthcare dataset from CSV.",
        "Remove duplicate rows and standardize column names.",
        "Validate patient ages and symptom-count consistency.",
        "Fill missing values and derive the age-group feature.",
        "Save the clean transformed dataset.",
        "Build the combined input string using age, gender, age group, and symptoms.",
        "Vectorize the input text using CountVectorizer.",
        "Train the Multinomial Naive Bayes classifier on the vectorized data.",
        "Accept user input from the Flask application.",
        "Transform the new input with the same vectorizer and predict probabilities.",
        "Sort the disease probabilities, normalize the top three, and present them in ranked order.",
        "Generate dashboard summaries and charts from the transformed dataset.",
    ]:
        add_number(document, item)

    add_heading(document, "4.11 DESIGN BENEFITS", level=2)
    benefits = [
        "The modular design supports maintenance because each file performs a single major responsibility.",
        "The chosen model is computationally light, making it appropriate for local demonstrations on standard systems.",
        "The dashboard adds transparency and helps reviewers understand data composition quickly.",
        "The CSV-to-database pathway shows readiness for future expansion without complicating the current project.",
        "The report-friendly architecture creates a natural mapping from implementation modules to academic chapters.",
    ]
    for item in benefits:
        add_bullet(document, item)


def chapter_five(document, stats):
    add_heading(document, "CHAPTER 5", level=1, centered=True, page_break=True)
    add_heading(document, "IMPLEMENTATION, ANALYTICS, AND RESULTS", level=1, centered=True)

    add_heading(document, "5.1 PROJECT FILE STRUCTURE", level=2)
    add_simple_table(
        document,
        [
            ("data.py", "Performs preprocessing and writes the transformed CSV"),
            ("model.py", "Builds features, trains the classifier, and predicts top diseases"),
            ("app.py", "Runs the Flask application and exposes user interaction"),
            ("datastore.py", "Uploads transformed data to MySQL as an optional extension"),
            ("Healthcare_data.csv", "Raw source dataset"),
            ("Healthcare_Transformed.csv", "Clean analytical dataset"),
            ("index.html / templates", "Web UI and dashboard rendering assets"),
        ],
    )
    add_paragraph(
        document,
        "The file organization reflects a good educational software structure. Data preparation, machine learning, "
        "and presentation are clearly separated, which improves readability and maintainability. This also makes it "
        "easy for reviewers to inspect one concern at a time without searching through a monolithic code file.",
    )

    add_heading(document, "5.2 PREPROCESSING IMPLEMENTATION", level=2)
    preprocessing = [
        "The preprocessing logic is implemented in data.py using Pandas. The script first reads the raw CSV file and "
        "checks duplicate rows. It then standardizes column names to a clean machine-friendly format. Because "
        "healthcare records can contain unrealistic ages or inconsistent symptom metadata, the script applies "
        "validation rules before the transformed file is saved.",
        "A calculated symptom count is obtained by splitting the symptom string and comparing the result against the "
        "provided symptom_count field. Rows that do not satisfy this consistency rule are removed. Numeric missing "
        "values are replaced with the median and categorical missing values with the mode. Such imputation keeps the "
        "pipeline simple and suitable for the project scale.",
        "The get_age_group function is then used to derive a new categorical field from the patient age. This field "
        "plays a meaningful role in later analytics because it groups records into interpretable segments rather than "
        "leaving age as an isolated number alone.",
    ]
    for text in preprocessing:
        add_paragraph(document, text)

    add_heading(document, "5.3 MODEL IMPLEMENTATION", level=2)
    model_details = [
        "In model.py, the transformed dataset is loaded and the symptom text is normalized by removing commas so that "
        "the token stream is cleaner. The age, gender, age group, and symptom text are concatenated into a single "
        "input string. This design is effective because it allows the vectorizer to capture patterns that jointly "
        "depend on demographics and symptoms.",
        "CountVectorizer transforms the combined input into a sparse count matrix, which is suitable for text-driven "
        "classification. The Multinomial Naive Bayes model is trained on this representation. For evaluation, the "
        "dataset is split into training and testing partitions with a fixed random state to make the experiment "
        "repeatable.",
        "The prediction function does more than return one class label. It computes class probabilities, sorts the "
        "disease-probability pairs, selects the top three classes, and normalizes those top values to a 100 percent "
        "scale. This creates a presentation-friendly output where users can view HIGH, MEDIUM, and LOW ranked "
        "disease options.",
    ]
    for text in model_details:
        add_paragraph(document, text)

    add_heading(document, "5.4 FLASK APPLICATION IMPLEMENTATION", level=2)
    flask_details = [
        "The Flask backend in app.py creates the connection between the prediction logic and the browser interface. "
        "The index route accepts both GET and POST requests. On form submission, the server reads the user input, "
        "derives the age group, and forwards the information to the prediction function. The ranked result is then "
        "rendered back into the HTML template.",
        "An additional API route returns the transformed dataset in JSON format. This route is useful because it "
        "shows that the same processed data can serve both the predictive module and data-driven front-end views. "
        "This small addition also demonstrates the project's extensibility, as future interfaces or mobile clients "
        "could consume the same endpoint.",
        "The Flask implementation remains intentionally lightweight. It is sufficient for demonstration, easy to run "
        "locally, and clean enough to explain during academic review. At the same time, it leaves space for future "
        "enhancements such as validation messages, authentication, persistent history, or integration with a "
        "database-backed user management module.",
    ]
    for text in flask_details:
        add_paragraph(document, text)

    add_heading(document, "5.5 DASHBOARD IMPLEMENTATION", level=2)
    dashboard = [
        "The project includes a separate dashboard page implemented with front-end technologies. The dashboard reads "
        "the transformed CSV file, parses the rows, computes counts and aggregate metrics, and renders multiple "
        "charts. This gives the project a broader analytical perspective beyond single-user prediction.",
        "The current dashboard shows summary cards for total records, top disease, and average age. It also renders "
        "charts for disease distribution, age-group distribution, gender split, disease versus age group, disease "
        "versus gender, and top-five disease comparison. These views are effective because they help reviewers connect "
        "the model output with the underlying dataset composition.",
        "By keeping the dashboard visually separate from the prediction form, the project supports more flexible "
        "demonstration. The user can first explain the dataset and then move to prediction, or start with patient "
        "input and later justify the results using the charts. This improves storytelling during presentation.",
    ]
    for text in dashboard:
        add_paragraph(document, text)

    add_heading(document, "5.6 DATASET ANALYTICS", level=2)
    add_three_column_table(
        document,
        ["Metric", "Value", "Interpretation"],
        [
            ("Total cleaned records", stats["clean_rows"], "Usable rows for training and dashboard analysis"),
            ("Age range", f"{stats['age_min']} - {stats['age_max']}", "Dataset spans multiple life stages"),
            ("Average age", stats["age_mean"], "Useful for dashboard summary"),
            ("Average symptoms per record", stats["symptom_mean"], "Shows symptom complexity per patient"),
            ("Top disease", list(stats["disease_counts"].keys())[0], "Most frequent label in the dataset"),
            ("Model accuracy", f"{stats['accuracy']}%", "Strong predictive performance on test split"),
        ],
    )
    add_paragraph(
        document,
        "Table 5.1 presents the main numerical observations from the cleaned dataset and the trained model. These "
        "values summarize the project at a glance and help the reviewer understand that the application is grounded "
        "in measurable outcomes rather than only interface design.",
    )

    add_heading(document, "5.7 DISEASE DISTRIBUTION ANALYSIS", level=2)
    disease_rows = [(name, count, f"{round((count / stats['clean_rows']) * 100, 2)}%") for name, count in stats["disease_counts"].items()]
    add_three_column_table(document, ["Disease", "Count", "Share"], disease_rows)
    disease_analysis = [
        "The disease distribution shows a moderately imbalanced but still useful dataset. Stroke, Diabetes, Heart "
        "Disease, Arthritis, and Anemia form the dominant set of classes. Influenza, Allergy, and Bronchitis occupy "
        "a middle range, while Thyroid Disorder appears with comparatively fewer examples.",
        "From a machine learning perspective, this distribution explains why a simple model can still perform well: "
        "the dataset contains strong recurring symptom patterns for the major disease classes. However, the low count "
        "for some categories also highlights why future versions should include more diverse samples.",
    ]
    for text in disease_analysis:
        add_paragraph(document, text)

    add_heading(document, "5.8 DEMOGRAPHIC ANALYSIS", level=2)
    demo_rows = []
    for key, value in stats["gender_counts"].items():
        demo_rows.append((f"Gender: {key}", value, f"{round((value / stats['clean_rows']) * 100, 2)}%"))
    for key, value in stats["age_group_counts"].items():
        demo_rows.append((f"Age Group: {key}", value, f"{round((value / stats['clean_rows']) * 100, 2)}%"))
    add_three_column_table(document, ["Category", "Count", "Share"], demo_rows)
    demo_analysis = [
        "The demographic distribution shows that female records are slightly higher than male records, while a small "
        "number of records are labeled under other gender categories. Adult records form the largest age-group "
        "segment, followed by seniors, children, and teenagers.",
        "These observations are useful because they reveal which population segments dominate the training data. In "
        "future expansions, the dataset could be balanced further so that minority groups receive even stronger "
        "representation in the learned model behavior.",
    ]
    for text in demo_analysis:
        add_paragraph(document, text)

    add_heading(document, "5.9 FREQUENT SYMPTOM ANALYSIS", level=2)
    symptom_rows = [(symptom, count, "Frequently occurring symptom token") for symptom, count in stats["top_symptoms"]]
    add_three_column_table(document, ["Symptom", "Count", "Comment"], symptom_rows)
    add_paragraph(
        document,
        "The frequent symptom list confirms that the dataset contains recurring descriptors that can be effectively "
        "modeled using a bag-of-words approach. High-frequency symptom tokens support stable classification because "
        "they create recognizable feature patterns across disease classes.",
    )

    add_heading(document, "5.10 MODEL EVALUATION", level=2)
    add_paragraph(
        document,
        f"The trained classifier achieved an overall test accuracy of {stats['accuracy']}%. This result is strong for "
        "the current dataset and validates the selected preprocessing and feature design strategy. The classification "
        "report below shows that precision, recall, and F1-score remain high across most disease classes.",
    )
    for line in stats["report"].strip().splitlines():
        if line.strip():
            add_paragraph(document, line, align="left", first_line=0, size=11, spacing_after=0, line_spacing=1.1)
    add_paragraph(
        document,
        "Although the current evaluation is highly positive, the result should still be interpreted in the context of "
        "the available dataset. Strong performance on a curated academic dataset does not automatically guarantee the "
        "same performance on broader clinical populations. This is an important point for responsible project "
        "presentation.",
    )

    add_heading(document, "5.11 RESULT INTERPRETATION", level=2)
    result_interp = [
        "The project demonstrates that a compact machine learning workflow can generate reliable disease predictions "
        "when paired with disciplined data cleaning. The quality of the transformed dataset directly supports the "
        "quality of the model. This is a valuable lesson because many student projects focus heavily on algorithms "
        "while underestimating the importance of preprocessing.",
        "The ranked output format is another strong result. Rather than giving a single rigid answer, the system "
        "reflects uncertainty through three ordered disease possibilities. This makes the interface more informative "
        "and aligns better with the symptom overlap commonly observed in healthcare-like problems.",
        "The dashboard strengthens result interpretation because it lets the reviewer see how disease counts, gender "
        "distribution, and age-group patterns shape the underlying dataset. This combination of prediction plus "
        "analytics makes the project easier to justify and more complete as an academic submission.",
    ]
    for text in result_interp:
        add_paragraph(document, text)

    add_heading(document, "5.12 SCREENSHOTS AND FIGURES", level=2)
    add_figure(document, IMAGE_DIR / "image1.jpg", "Fig. 5.1 Prediction module or supporting project visual")
    add_figure(document, IMAGE_DIR / "image2.jpg", "Fig. 5.2 Analytics dashboard or related presentation visual")
    add_figure(document, IMAGE_DIR / "image3.png", "Fig. 5.3 Additional system screen or chart snapshot")
    add_figure(document, IMAGE_DIR / "image4.png", "Fig. 5.4 Supplementary project visual or presentation asset")

    add_heading(document, "5.13 SAMPLE EXECUTION WALKTHROUGH", level=2)
    execution = [
        "When the application starts, Flask loads the trained model context and waits for user interaction. The home "
        "page presents fields for name, age, gender, and symptoms. Once the user enters the values and submits the "
        "form, the server route extracts the values from the request object and invokes the helper logic for age-group "
        "derivation.",
        "The derived age group is important because it enriches the feature space. A patient aged 11 is mapped to "
        "Child, while a patient aged 58 is mapped to Adult. This grouping gives the model another categorical cue "
        "that can differentiate diseases with similar symptoms but different demographic tendencies.",
        "The prediction function removes commas from the symptom string, constructs the combined feature text, "
        "transforms it using the already trained vectorizer, and predicts a full probability distribution over all "
        "supported diseases. The probabilities are sorted, the top three values are selected, and the results are "
        "returned as a dictionary containing HIGH, MEDIUM, and LOW risk labels.",
        "The rendered result page then presents the ranked outcomes in a presentation-friendly way. During a viva, "
        "this flow can be demonstrated live by entering different symptom combinations and explaining how the same "
        "data preparation and classification logic are reused for every request.",
    ]
    for text in execution:
        add_paragraph(document, text)

    add_heading(document, "5.14 API AND DATA SHARING VIEW", level=2)
    api_section = [
        "The API route defined in app.py exposes the transformed dataset in JSON format. This route shows that the "
        "project is not limited to server-side rendering alone. The same cleaned data can also be shared with "
        "front-end analytics components, external tools, or future application layers.",
        "From a software engineering perspective, this route is important because it separates data access from the "
        "HTML rendering process. If the project grows in future, a mobile app, a richer dashboard, or another "
        "frontend stack could consume the same JSON output without changing the core cleaning and modeling logic.",
        "This design pattern improves reusability and demonstrates that the project already contains the foundations "
        "of a service-oriented architecture, even though the current prototype remains intentionally lightweight.",
    ]
    for text in api_section:
        add_paragraph(document, text)

    add_heading(document, "5.15 RESULT QUALITY SUMMARY", level=2)
    quality = [
        "The high accuracy achieved by the model is supported by multiple aligned design decisions: validated input "
        "records, consistent symptom representation, a suitable vectorizer, and a classifier that matches sparse "
        "token-count features well.",
        "The result quality is also reflected in the usability of the interface. A technically accurate model has "
        "limited value if its output is difficult to interpret, but this project addresses that by normalizing the top "
        "three probabilities and giving them clear labels. This makes the result easier to discuss during project "
        "demonstration and easier for non-technical users to understand.",
        "Finally, the result quality is strengthened by the presence of the dashboard because aggregate evidence "
        "supports individual predictions. Together, the system's modules create a more complete academic artifact than "
        "a standalone script or isolated chart page.",
    ]
    for text in quality:
        add_paragraph(document, text)


def chapter_six(document):
    add_heading(document, "CHAPTER 6", level=1, centered=True, page_break=True)
    add_heading(document, "DISCUSSION, LIMITATIONS, AND FUTURE SCOPE", level=1, centered=True)

    add_heading(document, "6.1 DISCUSSION", level=2)
    discussion = [
        "The completed project demonstrates a healthy balance between software engineering, data analytics, and "
        "academic presentation. The preprocessing workflow ensures data quality. The machine learning pipeline remains "
        "simple enough to explain clearly. The Flask interface offers interactive use, and the dashboard provides "
        "supporting evidence through visual analytics.",
        "One of the strongest features of the work is that it does not treat prediction in isolation. The system also "
        "helps the user understand the dataset that drives those predictions. This combination is especially valuable "
        "in mini-capstone contexts where evaluation depends not only on code execution but also on clarity of thought, "
        "scope definition, and ability to communicate technical decisions.",
        "The project also demonstrates a practical principle in applied machine learning: a simple model with good "
        "data preparation can outperform more complicated approaches that are poorly structured or weakly explained. "
        "This lesson is worth emphasizing in the final submission because it reflects responsible engineering rather "
        "than only algorithmic ambition.",
    ]
    for text in discussion:
        add_paragraph(document, text)

    add_heading(document, "6.2 LIMITATIONS", level=2)
    limitations = [
        "The current dataset is limited in size and diversity. It is suitable for academic demonstration, but it does "
        "not represent the full complexity of real-world clinical records.",
        "The symptoms are modeled using a bag-of-words representation. While effective here, this approach does not "
        "capture deeper linguistic meaning, severity ordering, temporal progression, or contextual relationships.",
        "The model uses age, gender, age group, and symptoms only. It does not incorporate laboratory values, "
        "medication history, family history, comorbidities, or doctor notes.",
        "The application provides ranked suggestions but does not include confidence explanations beyond normalized "
        "probabilities. It should therefore be positioned as an educational support system rather than a clinical tool.",
        "The current deployment is local and does not include authentication, role management, audit logging, or "
        "secure persistent storage. Those are important for any serious healthcare-oriented deployment.",
    ]
    for item in limitations:
        add_bullet(document, item)

    add_heading(document, "6.3 FUTURE SCOPE", level=2)
    future_scope = [
        "Expand the dataset with additional disease categories and more balanced class representation.",
        "Compare Naive Bayes with Random Forest, Support Vector Machine, Logistic Regression, and transformer-based models.",
        "Introduce symptom severity, duration, and patient history as additional input features.",
        "Integrate the transformed dataset with a production-ready database and user authentication workflow.",
        "Add downloadable PDF reports, trend filters, and advanced dashboard drill-down views.",
        "Expose the prediction engine as a reusable API for mobile or institutional interfaces.",
        "Perform cross-validation and robustness analysis using multiple splits and external data samples.",
        "Add explainability features that highlight which input terms contributed most strongly to the predicted output.",
    ]
    for item in future_scope:
        add_bullet(document, item)

    add_heading(document, "6.4 ETHICAL AND PRACTICAL CONSIDERATIONS", level=2)
    ethics = [
        "Healthcare-related predictive systems must be presented carefully because users may confuse educational "
        "predictions with clinical diagnosis. This report therefore emphasizes that the application is a decision-"
        "support and learning prototype. It should never be used as a substitute for qualified medical evaluation.",
        "Another consideration involves privacy. Although the current dataset is used in an academic context, any "
        "future deployment would need strict safeguards for patient confidentiality, secure storage, role-based "
        "access, and compliance with relevant healthcare data standards.",
        "Bias and representational imbalance are additional concerns. If some diseases or demographic groups appear "
        "less often in the training data, prediction quality for those groups may differ. Recognizing such limits is "
        "part of responsible system design and should be highlighted during project presentation.",
    ]
    for text in ethics:
        add_paragraph(document, text)


def chapter_seven(document):
    add_heading(document, "CHAPTER 7", level=1, centered=True, page_break=True)
    add_heading(document, "CONCLUSION", level=1, centered=True)
    paragraphs = [
        "The Healthcare Data Analytics System developed in this project successfully transforms a structured patient "
        "dataset into both predictive and descriptive outputs. Through preprocessing, feature engineering, machine "
        "learning, and web integration, the project demonstrates how healthcare data can be converted into meaningful "
        "insight using an explainable and academically suitable approach.",
        "The project achieves multiple goals at once. It produces a cleaned transformed dataset, trains a disease "
        "prediction model with strong accuracy, provides ranked results through a Flask form, and visualizes the "
        "dataset through a dashboard. This integrated design gives the report more depth and makes the software easier "
        "to present in front of faculty reviewers.",
        "From an academic standpoint, the project is valuable because it shows disciplined engineering practice. The "
        "workflow is structured, the code responsibilities are separated, and the final documentation explains not "
        "only what was built but also why each design decision was taken. The expanded report therefore serves both "
        "as a submission document and as a presentation aid for viva and review.",
        "In summary, the work demonstrates that a practical healthcare analytics prototype can be built effectively "
        "with a strong focus on data quality, appropriate model selection, and user-friendly result presentation. It "
        "provides a solid foundation for future enhancement into a richer and more deployable healthcare support system.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)


def appendix_one(document):
    add_heading(document, "APPENDIX I", level=1, centered=True, page_break=True)
    add_heading(document, "SAMPLE INPUTS, OUTPUTS, AND EXPLANATORY NOTES", level=1, centered=True)
    paragraphs = [
        "The project accepts patient name, age, gender, and symptom descriptions through the web interface. The "
        "entered values are converted into an age group and then passed to the prediction engine. The engine returns "
        "three ranked disease outcomes labeled HIGH, MEDIUM, and LOW. This appendix documents the user interaction "
        "flow in descriptive form so that reviewers can quickly understand the behavior of the application.",
        "Sample use case 1: A user enters age 47, gender Male, and symptoms such as stiffness, swelling, and reduced "
        "mobility. Based on the transformed dataset patterns, the system is expected to prioritize disease classes "
        "related to arthritis-like conditions. The ranked probabilities help the user see the leading outcome along "
        "with secondary possibilities.",
        "Sample use case 2: A user enters age 57, gender Male, and symptoms such as fatigue, chest pain, and shortness "
        "of breath. In this context, the system may return heart-related disease categories near the top of the "
        "ranking because similar combinations are strongly represented in the dataset.",
        "Sample use case 3: A user enters age 11 with symptoms such as mucus, chest discomfort, and fatigue. The "
        "derived age group becomes Child, allowing the model to use an additional contextual feature while computing "
        "probabilities. This illustrates the value of feature engineering beyond raw symptom words alone.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)
    for item in [
        "Input validation note: the current academic prototype expects properly formatted values from the form.",
        "Prediction note: the output is a decision-support indication and not a substitute for medical diagnosis.",
        "Dashboard note: summary cards and charts are derived from the transformed dataset only.",
        "Presentation note: screenshot assets from the project presentation have been inserted to improve completeness.",
    ]:
        add_bullet(document, item)


def appendix_two(document):
    add_heading(document, "APPENDIX II", level=1, centered=True, page_break=True)
    add_heading(document, "CODE STRUCTURE AND FILE DESCRIPTION", level=1, centered=True)
    paragraphs = [
        "This appendix provides a narrative explanation of the major source files included in the project directory. "
        "The purpose is to help the reviewer connect the report content with the actual implementation files without "
        "needing to open every file during the viva.",
        "The file data.py focuses on data cleaning and transformation. It reads the raw dataset, removes duplicate "
        "rows, standardizes column names, validates age values, checks symptom-count consistency, handles missing "
        "values, derives the age group feature, and writes the transformed dataset to CSV.",
        "The file model.py focuses on training and prediction. It creates the input representation, vectorizes the "
        "combined text, trains the Multinomial Naive Bayes model, evaluates accuracy, and defines the reusable "
        "prediction function that ranks the top three diseases.",
        "The file app.py acts as the web server entry point. It creates the Flask application, receives user input, "
        "derives the age group, invokes the prediction function, renders the result, and also exposes an API route "
        "that returns dataset rows in JSON format.",
        "The file datastore.py demonstrates how the transformed CSV can be inserted into a MySQL database using "
        "SQLAlchemy. Even though the main project already works with CSV, this file shows an important direction for "
        "future expansion toward persistent storage and broader system integration.",
        "The HTML and front-end assets define the visual interface of the system. Together with the CSV, they enable "
        "summary cards and charts that support exploratory analysis and project explanation.",
    ]
    for text in paragraphs:
        add_paragraph(document, text)


def appendix_three(document):
    add_heading(document, "APPENDIX III", level=1, centered=True, page_break=True)
    add_heading(document, "VIVA PREPARATION NOTES", level=1, centered=True)
    intro = [
        "This appendix collects short explanatory points that can help during project review or viva. The purpose is "
        "to make the report more useful as a speaking aid in addition to being a written submission.",
        "A common viva question is why Naive Bayes was chosen over more advanced algorithms. A strong answer is that "
        "the project values explainability, fast execution, and good compatibility with sparse symptom text features. "
        "For the current dataset, the algorithm provides excellent performance while remaining easy to justify.",
        "Another likely question concerns the role of preprocessing. The correct explanation is that preprocessing is "
        "essential because it improves data consistency and directly affects the quality of model learning. Duplicate "
        "records, invalid ages, and inconsistent symptom counts can mislead the classifier if left untreated.",
        "Reviewers may also ask why the dashboard was included. The answer is that the dashboard gives descriptive "
        "analytics which complement the predictive output. It helps users understand class distribution, demographic "
        "spread, and other patterns that support the interpretation of the model results.",
    ]
    for text in intro:
        add_paragraph(document, text)
    qa_points = [
        "What is the input to the model? Age, gender, age group, and symptoms combined into one text string.",
        "Why is age group derived? It adds a meaningful demographic abstraction that helps pattern identification.",
        "What does the API route do? It exposes the transformed dataset in JSON form for reusable analytics access.",
        "Why are top three diseases shown instead of one? Because symptom overlap makes ranked output more informative.",
        "What is the main contribution of the project? Integration of preprocessing, prediction, dashboard analytics, and academic documentation.",
        "What is a major limitation? The dataset is academic and limited compared with real hospital-scale clinical data.",
        "How can the system be improved? Larger datasets, richer features, alternate algorithms, security, and deployment enhancements.",
    ]
    for item in qa_points:
        add_bullet(document, item)


def references(document):
    add_heading(document, "REFERENCES", level=1, centered=True, page_break=True)
    refs = [
        "Ian H. Witten, Eibe Frank, Mark A. Hall, and Christopher J. Pal, Data Mining: Practical Machine Learning Tools and Techniques.",
        "Tom M. Mitchell, Machine Learning.",
        "Pandas Documentation for data preprocessing and CSV handling.",
        "scikit-learn Documentation for CountVectorizer, train_test_split, and Multinomial Naive Bayes.",
        "Flask Documentation for Python web application development.",
        "SQLAlchemy Documentation for database connectivity and table creation.",
        "Chart.js Documentation for interactive chart rendering in browser environments.",
        "Project source files located in the current MCP workspace: data.py, model.py, app.py, datastore.py, Healthcare_data.csv, Healthcare_Transformed.csv, and HTML dashboard assets.",
    ]
    for ref in refs:
        add_number(document, ref)


def main():
    stats = build_stats()
    document = Document()
    set_document_language(document)
    set_page_layout(document)

    build_cover_page(document)
    build_certificate_page(document)
    build_acknowledgement(document)
    build_abstract(document, stats)
    build_contents(document)
    chapter_one(document, stats)
    chapter_two(document)
    chapter_three(document)
    chapter_four(document, stats)
    chapter_five(document, stats)
    chapter_six(document)
    chapter_seven(document)
    appendix_one(document)
    appendix_two(document)
    appendix_three(document)
    references(document)

    document.save(OUTPUT_FILE)
    print(f"Saved expanded report to: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
